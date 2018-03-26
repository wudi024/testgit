import os,re,sys
import win32com
from win32com.client import Dispatch, constants
from docx import Document

def parse_docx(file):
    d = Document(file)
    paras = d.paragraphs#段落
    tables = d.tables#表格

    print (file+'\t共有'+str(len(paras))+'个段落')
    print (file+'\t共有'+str(len(tables))+'个表格')
    try:
        f =open('D:/99bill.com/接口使用手册/输出文本2.txt', 'a', encoding='utf-8')
        f.write(file+'\t共有'+str(len(paras))+'个段落，'+str(len(tables))+'个表格\n\n')
        
        for i in range(len(paras)):
            if re.findall('^[A|B|C|M|P|S]+[0-9]+', paras[i].text):#接口格式：A|B|C|M|P|S开头+多个数字
                print (type(paras[i]))
                print ('是接口名称吧~~'+paras[i].text)
                f.write("第"+str(i)+"段的内容是："+paras[i].text+'\n')

        for i in range(len(tables)):
            
            print ('第'+str(i)+'个表格共有'+str(len(tables[i].rows))+'行')
            f.write('第'+str(i)+'个表格共有'+str(len(tables[i].rows))+'行\n')
            j=0
            t=tables[i]
            while(j<len(t.rows)):
                if re.findall('请求', t.cell(j,0).text):
                    break
                else:
                    j=j+1
            for row in range(j,len(t.rows)):
                print (type(tables[i].rows))
                if re.findall('请求', t.cell(row,0).text):
                    print (t.cell(row,0).text+'在第'+str(row+1)+'行')
                    f.write(t.cell(row,0).text+'在第'+str(row+1)+'行\n')
                    continue
                elif re.findall('响应', t.cell(row,0).text) :
                    print (t.cell(row,0).text+'在第'+str(row+1)+'行')
                    f.write(t.cell(row,0).text+'在第'+str(row+1)+'行\n')
                    continue
                elif str(t.cell(row,0).text).strip()!='':
                    reqParam = t.cell(row,0).text
                    reqType = t.cell(row,1).text
                    paraName = t.cell(row,2).text
                    print (reqParam.ljust(20)+ reqType.ljust(10)+paraName)
                    #f.write(reqParam.ljust(20)+ reqType.ljust(10)+'\t\t'+paraName+'\n')
                    f.write(reqParam+'\t\t'+ reqType+'\t\t'+paraName+'\n')

                
    finally:
        f.close()
        
if __name__ == "__main__":
    w = win32com.client.Dispatch('Word.Application')
    # 遍历文件
    PATH = "D:\\99bill.com\\接口使用手册" # windows文件路径
    doc_files = os.listdir(PATH)
    for doc in doc_files:
        if os.path.splitext(doc)[1] == '.docx':
            try:
                parse_docx(PATH+'\\'+doc)
            except Exception as e:
                print (e)
